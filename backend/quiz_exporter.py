"""
Quiz export functionality supporting multiple formats:
- PDF export with professional formatting
- Word document export
- JSON export for data interchange
- Moodle XML format for LMS integration
- Plain text format
"""

import json
import xml.etree.ElementTree as ET
from typing import List, Dict, Any, Optional
from io import BytesIO
import base64
from datetime import datetime

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("Warning: reportlab not installed. PDF export not available.")

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Word export not available.")

class QuizExporter:
    def __init__(self):
        self.supported_formats = ['json', 'txt', 'xml']
        
        if REPORTLAB_AVAILABLE:
            self.supported_formats.append('pdf')
        
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')

    def export_quiz(self, questions: List[Dict[str, Any]], 
                   format_type: str, 
                   title: str = "Generated Quiz",
                   metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Export quiz in specified format
        
        Args:
            questions: List of question dictionaries
            format_type: Export format ('pdf', 'docx', 'json', 'xml', 'txt')
            title: Quiz title
            metadata: Additional metadata (author, date, etc.)
            
        Returns:
            Dictionary with exported content and metadata
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}. Supported: {self.supported_formats}")
        
        if not metadata:
            metadata = {
                'generated_at': datetime.now().isoformat(),
                'total_questions': len(questions),
                'title': title
            }
        
        if format_type == 'json':
            return self._export_json(questions, title, metadata)
        elif format_type == 'txt':
            return self._export_txt(questions, title, metadata)
        elif format_type == 'xml':
            return self._export_moodle_xml(questions, title, metadata)
        elif format_type == 'pdf' and REPORTLAB_AVAILABLE:
            return self._export_pdf(questions, title, metadata)
        elif format_type == 'docx' and DOCX_AVAILABLE:
            return self._export_docx(questions, title, metadata)
        else:
            raise ValueError(f"Format {format_type} not available due to missing dependencies")

    def _export_json(self, questions: List[Dict], title: str, metadata: Dict) -> Dict[str, Any]:
        """Export quiz as JSON"""
        quiz_data = {
            'title': title,
            'metadata': metadata,
            'questions': questions
        }
        
        json_content = json.dumps(quiz_data, indent=2, ensure_ascii=False)
        
        return {
            'content': json_content,
            'content_type': 'application/json',
            'filename': f"{title.replace(' ', '_')}_quiz.json",
            'encoding': 'utf-8'
        }

    def _export_txt(self, questions: List[Dict], title: str, metadata: Dict) -> Dict[str, Any]:
        """Export quiz as plain text"""
        lines = []
        lines.append(f"Quiz Title: {title}")
        lines.append(f"Generated on: {metadata.get('generated_at', 'Unknown')}")
        lines.append(f"Total Questions: {metadata.get('total_questions', len(questions))}")
        lines.append("=" * 50)
        lines.append("")
        
        for i, question in enumerate(questions, 1):
            lines.append(f"Question {i}: {question.get('question', '')}")
            
            q_type = question.get('type', 'multiple_choice')
            
            if q_type == 'multiple_choice' and 'options' in question:
                for letter, option in question['options'].items():
                    marker = "* " if letter == question.get('correct_answer') else "  "
                    lines.append(f"{marker}{letter}. {option}")
            elif q_type == 'true_false':
                correct = question.get('correct_answer', 'True')
                lines.append(f"  True (*)  " if correct == 'True' else "  True")
                lines.append(f"  False (*)  " if correct == 'False' else "  False")
            elif q_type in ['fill_in_blank', 'short_answer']:
                lines.append(f"Answer: {question.get('correct_answer', '')}")
            
            if 'explanation' in question:
                lines.append(f"Explanation: {question['explanation']}")
            
            lines.append(f"Difficulty: {question.get('difficulty', 'medium')}")
            lines.append(f"Category: {question.get('category', 'general')}")
            lines.append("")
        
        content = "\n".join(lines)
        
        return {
            'content': content,
            'content_type': 'text/plain',
            'filename': f"{title.replace(' ', '_')}_quiz.txt",
            'encoding': 'utf-8'
        }

    def _export_moodle_xml(self, questions: List[Dict], title: str, metadata: Dict) -> Dict[str, Any]:
        """Export quiz in Moodle XML format"""
        quiz = ET.Element('quiz')
        
        for i, question in enumerate(questions, 1):
            q_element = ET.SubElement(quiz, 'question')
            q_type = question.get('type', 'multichoice')
            
            # Map our types to Moodle types
            moodle_types = {
                'multiple_choice': 'multichoice',
                'true_false': 'truefalse',
                'fill_in_blank': 'shortanswer',
                'short_answer': 'shortanswer',
                'matching': 'matching'
            }
            
            q_element.set('type', moodle_types.get(q_type, 'multichoice'))
            
            # Question name
            name = ET.SubElement(q_element, 'name')
            name_text = ET.SubElement(name, 'text')
            name_text.text = f"{title} - Question {i}"
            
            # Question text
            questiontext = ET.SubElement(q_element, 'questiontext')
            questiontext.set('format', 'html')
            qt_text = ET.SubElement(questiontext, 'text')
            qt_text.text = f"<![CDATA[{question.get('question', '')}]]>"
            
            # Default grade
            default_grade = ET.SubElement(q_element, 'defaultgrade')
            default_grade.text = str(question.get('points', 1))
            
            # Question-specific elements
            if q_type == 'multiple_choice':
                self._add_moodle_multichoice_elements(q_element, question)
            elif q_type == 'true_false':
                self._add_moodle_truefalse_elements(q_element, question)
            elif q_type in ['fill_in_blank', 'short_answer']:
                self._add_moodle_shortanswer_elements(q_element, question)
            
            # General feedback
            if 'explanation' in question:
                feedback = ET.SubElement(q_element, 'generalfeedback')
                feedback.set('format', 'html')
                fb_text = ET.SubElement(feedback, 'text')
                fb_text.text = question['explanation']
        
        # Convert to string
        xml_string = ET.tostring(quiz, encoding='unicode', method='xml')
        formatted_xml = '<?xml version="1.0" encoding="UTF-8"?>\n' + xml_string
        
        return {
            'content': formatted_xml,
            'content_type': 'application/xml',
            'filename': f"{title.replace(' ', '_')}_moodle.xml",
            'encoding': 'utf-8'
        }

    def _add_moodle_multichoice_elements(self, q_element: ET.Element, question: Dict):
        """Add Moodle-specific elements for multiple choice questions"""
        # Single answer (not multiple)
        single = ET.SubElement(q_element, 'single')
        single.text = 'true'
        
        # Shuffle answers
        shuffle_answers = ET.SubElement(q_element, 'shuffleanswers')
        shuffle_answers.text = 'true'
        
        # Answer options
        if 'options' in question:
            correct_answer = question.get('correct_answer', 'A')
            for letter, option_text in question['options'].items():
                answer = ET.SubElement(q_element, 'answer')
                fraction = '100' if letter == correct_answer else '0'
                answer.set('fraction', fraction)
                answer.set('format', 'html')
                
                a_text = ET.SubElement(answer, 'text')
                a_text.text = f"<![CDATA[{option_text}]]>"

    def _add_moodle_truefalse_elements(self, q_element: ET.Element, question: Dict):
        """Add Moodle-specific elements for true/false questions"""
        correct = question.get('correct_answer', 'True') == 'True'
        
        # True answer
        true_answer = ET.SubElement(q_element, 'answer')
        true_answer.set('fraction', '100' if correct else '0')
        true_answer.set('format', 'moodle_auto_format')
        true_text = ET.SubElement(true_answer, 'text')
        true_text.text = 'true'
        
        # False answer
        false_answer = ET.SubElement(q_element, 'answer')
        false_answer.set('fraction', '0' if correct else '100')
        false_answer.set('format', 'moodle_auto_format')
        false_text = ET.SubElement(false_answer, 'text')
        false_text.text = 'false'

    def _add_moodle_shortanswer_elements(self, q_element: ET.Element, question: Dict):
        """Add Moodle-specific elements for short answer questions"""
        answer = ET.SubElement(q_element, 'answer')
        answer.set('fraction', '100')
        answer.set('format', 'moodle_auto_format')
        
        a_text = ET.SubElement(answer, 'text')
        a_text.text = question.get('correct_answer', '')

    def _export_pdf(self, questions: List[Dict], title: str, metadata: Dict) -> Dict[str, Any]:
        """Export quiz as PDF using ReportLab"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'QuizTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        question_style = ParagraphStyle(
            'QuestionText',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            leftIndent=0
        )
        
        option_style = ParagraphStyle(
            'OptionText',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            spaceAfter=5
        )
        
        story = []
        
        # Title and metadata
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Metadata table
        metadata_data = [
            ['Generated:', metadata.get('generated_at', 'Unknown')[:10]],
            ['Questions:', str(metadata.get('total_questions', len(questions)))]
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 30))
        
        # Questions
        for i, question in enumerate(questions, 1):
            # Question number and text
            story.append(Paragraph(
                f"<b>Question {i}:</b> {question.get('question', '')}", 
                question_style
            ))
            
            q_type = question.get('type', 'multiple_choice')
            
            if q_type == 'multiple_choice' and 'options' in question:
                for letter, option in question['options'].items():
                    story.append(Paragraph(f"{letter}. {option}", option_style))
            elif q_type == 'true_false':
                story.append(Paragraph("A. True", option_style))
                story.append(Paragraph("B. False", option_style))
            elif q_type in ['fill_in_blank', 'short_answer']:
                story.append(Paragraph("Answer: ___________________", option_style))
            
            # Answer key and metadata
            answer_info = f"Correct Answer: {question.get('correct_answer', 'N/A')} | "
            answer_info += f"Difficulty: {question.get('difficulty', 'medium')} | "
            answer_info += f"Category: {question.get('category', 'general')}"
            
            story.append(Paragraph(
                f"<i>{answer_info}</i>", 
                ParagraphStyle('AnswerInfo', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
            ))
            
            story.append(Spacer(1, 20))
        
        doc.build(story)
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return {
            'content': base64.b64encode(pdf_content).decode('utf-8'),
            'content_type': 'application/pdf',
            'filename': f"{title.replace(' ', '_')}_quiz.pdf",
            'encoding': 'base64'
        }

    def _export_docx(self, questions: List[Dict], title: str, metadata: Dict) -> Dict[str, Any]:
        """Export quiz as Word document"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {metadata.get('generated_at', 'Unknown')[:10]}")
        doc.add_paragraph(f"Total Questions: {metadata.get('total_questions', len(questions))}")
        doc.add_paragraph()  # Empty line
        
        # Questions
        for i, question in enumerate(questions, 1):
            # Question
            q_para = doc.add_paragraph()
            q_para.add_run(f"Question {i}: ").bold = True
            q_para.add_run(question.get('question', ''))
            
            q_type = question.get('type', 'multiple_choice')
            
            if q_type == 'multiple_choice' and 'options' in question:
                for letter, option in question['options'].items():
                    doc.add_paragraph(f"{letter}. {option}", style='List Number')
            elif q_type == 'true_false':
                doc.add_paragraph("A. True", style='List Number')
                doc.add_paragraph("B. False", style='List Number')
            elif q_type in ['fill_in_blank', 'short_answer']:
                doc.add_paragraph("Answer: ___________________")
            
            # Answer and metadata
            answer_para = doc.add_paragraph()
            answer_para.add_run("Answer: ").italic = True
            answer_para.add_run(str(question.get('correct_answer', 'N/A'))).italic = True
            
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Difficulty: {question.get('difficulty', 'medium')} | ").italic = True
            meta_para.add_run(f"Category: {question.get('category', 'general')}").italic = True
            
            doc.add_paragraph()  # Empty line between questions
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        docx_content = buffer.getvalue()
        buffer.close()
        
        return {
            'content': base64.b64encode(docx_content).decode('utf-8'),
            'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'filename': f"{title.replace(' ', '_')}_quiz.docx",
            'encoding': 'base64'
        }

# Factory function
def create_quiz_exporter() -> QuizExporter:
    """Create quiz exporter instance"""
    return QuizExporter()